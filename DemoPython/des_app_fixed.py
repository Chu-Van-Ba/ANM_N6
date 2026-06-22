import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import random
import os

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

IP = [58,50,42,34,26,18,10,2,
      60,52,44,36,28,20,12,4,
      62,54,46,38,30,22,14,6,
      64,56,48,40,32,24,16,8,
      57,49,41,33,25,17,9,1,
      59,51,43,35,27,19,11,3,
      61,53,45,37,29,21,13,5,
      63,55,47,39,31,23,15,7]


IP_INV = [40,8,48,16,56,24,64,32,
          39,7,47,15,55,23,63,31,
          38,6,46,14,54,22,62,30,
          37,5,45,13,53,21,61,29,
          36,4,44,12,52,20,60,28,
          35,3,43,11,51,19,59,27,
          34,2,42,10,50,18,58,26,
          33,1,41,9,49,17,57,25]


E = [32,1,2,3,4,5,
     4,5,6,7,8,9,
     8,9,10,11,12,13,
     12,13,14,15,16,17,
     16,17,18,19,20,21,
     20,21,22,23,24,25,
     24,25,26,27,28,29,
     28,29,30,31,32,1]


P = [16,7,20,21,29,12,28,17,
     1,15,23,26,5,18,31,10,
     2,8,24,14,32,27,3,9,
     19,13,30,6,22,11,4,25]


S = [
    # S1
    [[14,4,13,1,2,15,11,8,3,10,6,12,5,9,0,7],
     [0,15,7,4,14,2,13,1,10,6,12,11,9,5,3,8],
     [4,1,14,8,13,6,2,11,15,12,9,7,3,10,5,0],
     [15,12,8,2,4,9,1,7,5,11,3,14,10,0,6,13]],
    # S2
    [[15,1,8,14,6,11,3,4,9,7,2,13,12,0,5,10],
     [3,13,4,7,15,2,8,14,12,0,1,10,6,9,11,5],
     [0,14,7,11,10,4,13,1,5,8,12,6,9,3,2,15],
     [13,8,10,1,3,15,4,2,11,6,7,12,0,5,14,9]],
    # S3
    [[10,0,9,14,6,3,15,5,1,13,12,7,11,4,2,8],
     [13,7,0,9,3,4,6,10,2,8,5,14,12,11,15,1],
     [13,6,4,9,8,15,3,0,11,1,2,12,5,10,14,7],
     [1,10,13,0,6,9,8,7,4,15,14,3,11,5,2,12]],
    # S4
    [[7,13,14,3,0,6,9,10,1,2,8,5,11,12,4,15],
     [13,8,11,5,6,15,0,3,4,7,2,12,1,10,14,9],
     [10,6,9,0,12,11,7,13,15,1,3,14,5,2,8,4],
     [3,15,0,6,10,1,13,8,9,4,5,11,12,7,2,14]],
    # S5
    [[2,12,4,1,7,10,11,6,8,5,3,15,13,0,14,9],
     [14,11,2,12,4,7,13,1,5,0,15,10,3,9,8,6],
     [4,2,1,11,10,13,7,8,15,9,12,5,6,3,0,14],
     [11,8,12,7,1,14,2,13,6,15,0,9,10,4,5,3]],
    # S6
    [[12,1,10,15,9,2,6,8,0,13,3,4,14,7,5,11],
     [10,15,4,2,7,12,9,5,6,1,13,14,0,11,3,8],
     [9,14,15,5,2,8,12,3,7,0,4,10,1,13,11,6],
     [4,3,2,12,9,5,15,10,11,14,1,7,6,0,8,13]],
    # S7
    [[4,11,2,14,15,0,8,13,3,12,9,7,5,10,6,1],
     [13,0,11,7,4,9,1,10,14,3,5,12,2,15,8,6],
     [1,4,11,13,12,3,7,14,10,15,6,8,0,5,9,2],
     [6,11,13,8,1,4,10,7,9,5,0,15,14,2,3,12]],
    # S8
    [[13,2,8,4,6,15,11,1,10,9,3,14,5,0,12,7],
     [1,15,13,8,10,3,7,4,12,5,6,11,0,14,9,2],
     [7,11,4,1,9,12,14,2,0,6,10,13,15,3,5,8],
     [2,1,14,7,4,10,8,13,15,12,9,0,3,5,6,11]]
]


PC1 = [57,49,41,33,25,17,9,
       1,58,50,42,34,26,18,
       10,2,59,51,43,35,27,
       19,11,3,60,52,44,36,
       63,55,47,39,31,23,15,
       7,62,54,46,38,30,22,
       14,6,61,53,45,37,29,
       21,13,5,28,20,12,4]


PC2 = [14,17,11,24,1,5,3,28,
       15,6,21,10,23,19,12,4,
       26,8,16,7,27,20,13,2,
       41,52,31,37,47,55,30,40,
       51,45,33,48,44,49,39,56,
       34,53,46,42,50,36,29,32]


SHIFTS = [1,1,2,2,2,2,2,2,1,2,2,2,2,2,2,1]


def permute(bits, table):
    return [bits[t-1] for t in table]

def left_shift(bits, n):
    return bits[n:] + bits[:n]

def xor(a, b):
    return [x ^ y for x, y in zip(a, b)]

def int_to_bits(n, length):
    return [int(b) for b in bin(n)[2:].zfill(length)]

def bits_to_int(bits):
    result = 0
    for b in bits:
        result = (result << 1) | b
    return result

def bits_to_hex(bits):
    n = bits_to_int(bits)
    return hex(n)[2:].upper().zfill(len(bits)//4)

def hex_to_bits(h, length=None):
    n = int(h, 16)
    if length is None:
        length = len(h) * 4
    return int_to_bits(n, length)

def str_to_bits(s):
    data = s.encode("utf-8")   # chuyển chuỗi sang bytes UTF-8
    bits = []
    for byte in data:
        bits += int_to_bits(byte, 8)
    return bits

def bits_to_str(bits):
    data = bytearray()

    for i in range(0, len(bits), 8):
        byte = bits_to_int(bits[i:i+8])
        data.append(byte)

    return data.decode("utf-8", errors="replace")

def generate_subkeys(key_bits):
    """Generate 16 subkeys from 64-bit key"""
    cd = permute(key_bits, PC1)     # Áp dụng PC1 → 56 bit
    c, d = cd[:28], cd[28:]         # Chia C₀ (28 bit) và D₀ (28 bit)
    subkeys = []
    for shift in SHIFTS:            # Duyệt 16 vòng
        c = left_shift(c, shift)    # Dịch vòng trái Cᵢ
        d = left_shift(d, shift)    # Dịch vòng trái Dᵢ
        subkey = permute(c + d, PC2)# Áp dụng PC2 → 48 bit
        subkeys.append(subkey)
    return subkeys

def f_function(right, subkey):
    """Feistel function"""
    expanded = permute(right, E)          # Mở rộng E: 32 → 48 bit
    xored = xor(expanded, subkey)         # XOR với khóa con 48 bit
    sbox_out = []
    for i in range(8):
        block = xored[i*6:(i+1)*6]        # Tách thành 8 khối 6 bit
        row = (block[0] << 1) | block[5]  # Hàng: b₁b₆
        col = bits_to_int(block[1:5])     # Cột: b₂b₃b₄b₅
        val = S[i][row][col]              # Tra S-box
        sbox_out += int_to_bits(val, 4)   # Kết quả 4 bit
    return permute(sbox_out, P)           # Hoán vị P → 32 bit

def des_block(block_bits, subkeys):
    """DES on a single 64-bit block"""
    ip = permute(block_bits, IP)        # Hoán vị IP
    left, right = ip[:32], ip[32:]      # Chia L₀, R₀
    for subkey in subkeys:              # 16 vòng Feistel
        new_right = xor(left, f_function(right, subkey))
        left = right
        right = new_right
    combined = right + left             # Ghép R₁₆L₁₆ (đổi trái-phải)
    return permute(combined, IP_INV)    # Hoán vị IP⁻¹

def pad_bits(bits):
    """PKCS#7-style padding to multiple of 64 bits"""
    remainder = len(bits) % 64
    if remainder == 0:
        pad_bytes = 8
    else:
        pad_bytes = (64 - remainder) // 8
    pad_bit = int_to_bits(pad_bytes, 8)
    for _ in range(pad_bytes):
        bits += pad_bit
    return bits

def unpad_bits(bits):
    """Remove PKCS#7 padding"""
    pad_bytes = bits_to_int(bits[-8:])
    if pad_bytes < 1 or pad_bytes > 8:
        return bits
    return bits[:-(pad_bytes * 8)]

def des_encrypt(plaintext, key_hex):
    """Encrypt plaintext string with key (hex string)"""
    key_bits = hex_to_bits(key_hex, 64)
    subkeys = generate_subkeys(key_bits)
    plain_bits = pad_bits(str_to_bits(plaintext))
    cipher_bits = []
    for i in range(0, len(plain_bits), 64):
        block = plain_bits[i:i+64]
        cipher_bits += des_block(block, subkeys)
    return bits_to_hex(cipher_bits)

def des_decrypt(ciphertext_hex, key_hex):
    """Decrypt hex ciphertext with key (hex string)"""
    key_bits = hex_to_bits(key_hex, 64)
    subkeys = generate_subkeys(key_bits)[::-1]  
    cipher_bits = hex_to_bits(ciphertext_hex)
    plain_bits = []
    for i in range(0, len(cipher_bits), 64):
        block = cipher_bits[i:i+64]
        plain_bits += des_block(block, subkeys)
    plain_bits = unpad_bits(plain_bits)
    return bits_to_str(plain_bits)

def generate_random_key():
    """Generate a random 64-bit key (16 hex characters)"""
    return ''.join(random.choice('0123456789ABCDEF') for _ in range(16))

def is_valid_hex(h, length=16):
    try:
        if len(h) != length:
            return False
        int(h, 16)
        return True
    except:
        return False

def read_docx_file(path):
    """Đọc nội dung văn bản từ file .docx"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("Chưa cài thư viện python-docx! Chạy: pip install python-docx")
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def write_docx_file(path, text):
    """Ghi nội dung văn bản ra file .docx"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("Chưa cài thư viện python-docx! Chạy: pip install python-docx")
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)

def read_pdf_file(path):
    """Đọc nội dung văn bản từ file .pdf"""
    if not PDF_AVAILABLE:
        raise RuntimeError("Chưa cài thư viện PyPDF2! Chạy: pip install PyPDF2")
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


BG        = "#B9D3EE"   # Nền thạch anh xanh nhạt 
SURFACE   = "#FFFFFF"   # Vùng panel màu trắng 
BORDER    = "#D4D4D4"   # Đường viền panel mảnh mai tinh tế
ACCENT1   = "#4A56E2"   # Tím Royal dải màu panel khóa
ACCENT2   = "#0B72D9"   # Xanh dương thẫm cho dải màu panel mã hóa
ACCENT3   = "#1A9E6B"   # Xanh lá thẫm cho dải màu panel giải mã
FG        = "#1E293B"   # Chữ chính xám đen 
FG_DIM    = "#64748B"   # Chữ mô tả xám nhạt
FG_MONO   = "#0F52BA"   # Màu chữ Monospace (Hex Output)

TEXT_BG   = "#C8DCF1"   # Nền các ô nhập liệu màu xanh nước biển 
TEXT_FG   = "#000000"   
MONO_BG   = "#E6F2FF"   
BTN_FG    = "#FFFFFF"


class DESApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Hệ mật mã DES")
        self.root.configure(bg=BG)
        self.root.resizable(True, True)

        self.key_mode = tk.StringVar(value="manual")
        self.last_plaintext = ""
        self.last_ciphertext_hex = ""
        self.last_key_used = ""

        self._build_ui()
        self.root.minsize(1024, 660)
        self.root.geometry("1200x740")

    def _build_ui(self):
        # Tiêu đề chính ứng dụng cân đối
        hdr = tk.Frame(self.root, bg=BG, pady=0)
        hdr.pack(fill=tk.X)

        top_strip = tk.Frame(hdr, bg=ACCENT1, height=3)
        top_strip.pack(fill=tk.X)

        title_row = tk.Frame(hdr, bg=BG, pady=16, padx=24)
        title_row.pack(fill=tk.X)

        tk.Label(title_row, text="DES",
                 font=("Courier New", 22, "bold"),
                 fg=ACCENT1, bg=BG).pack(side=tk.LEFT)
        tk.Label(title_row,
                 text="  Tạo khóa, Mã hóa & Giải mã dữ liệu",
                 font=("Segoe UI", 10, "bold"), fg=FG_DIM, bg=BG).pack(side=tk.LEFT, pady=6)

        sep = tk.Frame(hdr, bg=BORDER, height=1)
        sep.pack(fill=tk.X)

        # Container chia 3 cột chính
        main = tk.Frame(self.root, bg=BG, padx=16, pady=12)
        main.pack(fill=tk.BOTH, expand=True)

        main.columnconfigure(0, weight=1, minsize=260)
        main.columnconfigure(1, weight=1, minsize=260)
        main.columnconfigure(2, weight=1, minsize=260)
        main.rowconfigure(0, weight=1)

        self._build_key_panel(main)
        self._build_encrypt_panel(main)
        self._build_decrypt_panel(main)

        status = tk.Frame(self.root, bg=SURFACE, pady=6)
        status.pack(fill=tk.X, side=tk.BOTTOM)
        tk.Label(status,
                 text="  DES — 16 vòng Feistel  |  Khóa: 56-bit  |  Block: 64-bit  ",
                 font=("Courier New", 8, "bold"), fg=FG_DIM, bg=SURFACE).pack(side=tk.LEFT)

    def _build_key_panel(self, parent):
        frame = self._panel(parent, "[ 01 ]  TẠO KHÓA BÍ MẬT", ACCENT1, col=0)

        self._label(frame, "Chế độ nhập khóa:")
        rb_frame = tk.Frame(frame, bg=SURFACE)
        rb_frame.pack(anchor="w", pady=(4, 12)) 
        for txt, val in [("Tạo khóa thủ công", "manual"), ("Tạo khóa tự động", "auto")]:
            tk.Radiobutton(rb_frame, text=txt, variable=self.key_mode, value=val,
                           bg=SURFACE, fg=FG, selectcolor=SURFACE,
                           activebackground=SURFACE, activeforeground=FG,
                           font=("Segoe UI", 9),
                           command=self._on_key_mode_change).pack(side=tk.LEFT, padx=(0, 16))

        self._label(frame, "Khóa (16 ký tự HEX):")
        self.key_entry = tk.Text(frame, height=2,
                                 font=("Courier New", 11, "bold"),
                                 bg=MONO_BG, fg=FG_MONO,
                                 insertbackground=FG_MONO,
                                 relief=tk.FLAT, bd=0,
                                 wrap=tk.WORD,
                                 highlightthickness=1,
                                 highlightbackground=ACCENT1,
                                 highlightcolor=ACCENT1)
        self.key_entry.pack(fill=tk.X, pady=(4, 14)) 

        bf_key = tk.Frame(frame, bg=SURFACE)
        bf_key.pack(fill=tk.X, pady=(0, 12))
        self._btn(bf_key, "📂 Mở File", self._open_key_file,
                  bg="#065F46", hover="#047857", side=tk.LEFT)
        self._btn(bf_key, "💾 Lưu khóa", self._save_key_file,
                  bg="#374151", hover="#4B5563", side=tk.LEFT)
        self._btn(bf_key, "🗑️ Xóa", self._reset_key,
                  bg="#7F1D1D", hover="#991B1B", side=tk.LEFT)

        self._btn(frame, "➡️ Chuyển khóa sang Giải mã", self._copy_key_to_decrypt,
                  bg="#1D4ED8", hover="#2563EB")

        tk.Label(frame, text="Khóa dùng chung cho cả mã hóa và giải mã",
                 bg=SURFACE, fg=FG_DIM, font=("Segoe UI", 8)).pack(pady=(8, 0))
        # tk.Label(frame, text="Khóa dùng chung cho cả mã hóa và giải mã",
        #          bg=SURFACE, fg=FG_DIM, font=("Segoe UI", 9)).pack(side=tk.BOTTOM, pady=(15, 0))

    def _build_encrypt_panel(self, parent):
        frame = self._panel(parent, "[ 02 ]  MÃ HÓA", ACCENT2, col=1)

        self._label(frame, "Bản rõ (Plaintext):")
        self.plain_text = tk.Text(frame, height=6,
                                  font=("Segoe UI", 10),
                                  bg=TEXT_BG, fg=TEXT_FG,
                                  insertbackground=TEXT_FG,
                                  relief=tk.FLAT, bd=0,
                                  wrap=tk.WORD,
                                  highlightthickness=1,
                                  highlightbackground=BORDER,
                                  highlightcolor=ACCENT2)
        self.plain_text.pack(fill=tk.X, pady=(4, 14))

        bf = tk.Frame(frame, bg=SURFACE)
        bf.pack(fill=tk.X, pady=(0, 12))
        self._btn(bf, "📂 Mở File", self._open_plain_file,
                  bg="#065F46", hover="#047857", side=tk.LEFT)
        self._btn(bf, "💾 Lưu bản rõ", self._save_plain_input_file,
                  bg="#374151", hover="#4B5563", side=tk.LEFT)
        self._btn(bf, "🗑️ Xóa", lambda: self.plain_text.delete("1.0", tk.END),
                  bg="#7F1D1D", hover="#991B1B", side=tk.LEFT)

        self._btn(frame, "🔒 Mã hóa", self._encrypt,
                  bg=ACCENT2, hover="#0284C7")

        self._label(frame, "Bản mã (Ciphertext Output):", pady=(16, 0))
        self.cipher_text = tk.Text(frame, height=5,
                                   font=("Courier New", 9),
                                   bg=MONO_BG, fg=FG_MONO,
                                   relief=tk.FLAT, bd=0,
                                   wrap=tk.WORD,
                                   highlightthickness=1,
                                   highlightbackground=BORDER,
                                   highlightcolor=ACCENT2,
                                   state=tk.DISABLED)
        self.cipher_text.pack(fill=tk.X, pady=(4, 14))

        bf2 = tk.Frame(frame, bg=SURFACE)
        bf2.pack(fill=tk.X, pady=(0, 4))
        self._btn(bf2, "➡️ Chuyển sang Giải mã", self._transfer_cipher,
                  bg="#1E40AF", hover="#1D4ED8", side=tk.LEFT)
        self._btn(bf2, "💾 Lưu bản mã", self._save_cipher_file,
                  bg="#374151", hover="#4B5563", side=tk.LEFT)

    def _build_decrypt_panel(self, parent):
        frame = self._panel(parent, "[ 03 ]  GIẢI MÃ", ACCENT3, col=2)

        self._label(frame, "Bản mã (Ciphertext):")
        self.dec_input = tk.Text(frame, height=4,
                                 font=("Courier New", 9),
                                 bg=MONO_BG, fg=FG_MONO,
                                 insertbackground=FG_MONO,
                                 relief=tk.FLAT, bd=0,
                                 wrap=tk.WORD,
                                 highlightthickness=1,
                                 highlightbackground=BORDER,
                                 highlightcolor=ACCENT3)
        self.dec_input.pack(fill=tk.X, pady=(4, 14))

        bf = tk.Frame(frame, bg=SURFACE)
        bf.pack(fill=tk.X, pady=(0, 12))
        self._btn(bf, "📂 Mở File", self._open_cipher_file,
                  bg="#065F46", hover="#047857", side=tk.LEFT)
        self._btn(bf, "🗑️ Xóa", lambda: self.dec_input.delete("1.0", tk.END),
                  bg="#7F1D1D", hover="#991B1B", side=tk.LEFT)

        self._btn(frame, "🔓 Giải mã", self._decrypt,
                  bg=ACCENT3, hover="#059669")

        self._label(frame, "Khóa bí mật (16 ký tự HEX):", pady=(16, 0))
        self.dec_key_display = tk.Text(frame, height=2,
                                       font=("Courier New", 11, "bold"),
                                       bg=MONO_BG, fg=FG_MONO,
                                       insertbackground=FG_MONO,
                                       relief=tk.FLAT, bd=0,
                                       wrap=tk.WORD,
                                       highlightthickness=1,
                                       highlightbackground=ACCENT1,
                                       highlightcolor=ACCENT1)
        self.dec_key_display.pack(fill=tk.X, pady=(4, 14))

        self._label(frame, "Bản rõ (Plaintext Output):", pady=(16, 0))
        self.dec_output = tk.Text(frame, height=4,
                                  font=("Segoe UI", 10),
                                  bg=TEXT_BG, fg=TEXT_FG,
                                  relief=tk.FLAT, bd=0,
                                  wrap=tk.WORD,
                                  highlightthickness=1,
                                  highlightbackground=BORDER,
                                  highlightcolor=ACCENT3,
                                  state=tk.DISABLED)
        self.dec_output.pack(fill=tk.X, pady=(4, 14))

        bf2 = tk.Frame(frame, bg=SURFACE)
        bf2.pack(fill=tk.X, pady=(0, 4))
        self._btn(bf2, "💾 Lưu bản rõ", self._save_plain_file,
                  bg="#374151", hover="#4B5563", side=tk.LEFT)
        self._btn(bf2, "✅ Kiểm tra kết quả", self._verify_result,
                  bg="#5B21B6", hover="#6D28D9", side=tk.LEFT)

    def _panel(self, parent, title, accent, col):
        outer = tk.Frame(parent, bg=SURFACE,
                         highlightthickness=2,
                         highlightbackground=BORDER,
                         highlightcolor=BORDER)
        outer.grid(row=0, column=col, sticky="nsew", padx=8, pady=8) 

        bar = tk.Frame(outer, bg=accent, height=4) 
        bar.pack(fill=tk.X)

        title_row = tk.Frame(outer, bg=SURFACE, pady=10, padx=14)
        title_row.pack(fill=tk.X)
        tk.Label(title_row, text=title,
                 font=("Courier New", 11, "bold"),
                 bg=SURFACE, fg=accent).pack(anchor="w")

        sep = tk.Frame(outer, bg=BORDER, height=1)
        sep.pack(fill=tk.X, padx=0)

        inner = tk.Frame(outer, bg=SURFACE, padx=14, pady=12)
        inner.pack(fill=tk.BOTH, expand=True)
        return inner

    def _label(self, parent, text, pady=(0, 0)):
        tk.Label(parent, text=text,
                 font=("Segoe UI", 9, "bold"),
                 bg=SURFACE, fg=FG_DIM).pack(anchor="w", pady=pady)

    def _btn(self, parent, text, command,
             bg="#374151", hover="#4B5563", fg=BTN_FG, side=None):
        b = tk.Button(parent, text=text, command=command,
                      bg=bg, fg=fg,
                      font=("Segoe UI", 9, "bold"),
                      relief=tk.FLAT, bd=0,
                      padx=16, pady=8,
                      cursor="hand2",
                      activebackground=hover,
                      activeforeground=fg,
                      highlightbackground=SURFACE,   
                      highlightthickness=0)
        if side:
            b.pack(side=side, padx=(0, 6), fill=tk.X, expand=True)
        else:
            b.pack(fill=tk.X, pady=(0, 8))
        return b

    def _get_key(self):
        return self.key_entry.get("1.0", tk.END).strip().upper()

    def _get_dec_key(self):
        return self.dec_key_display.get("1.0", tk.END).strip().upper()

    def _set_text(self, widget, text):
        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert("1.0", text)

    def _set_readonly(self, widget, text):
        widget.config(state=tk.NORMAL)
        widget.delete("1.0", tk.END)
        widget.insert("1.0", text)
        widget.config(state=tk.DISABLED)

    def _on_key_mode_change(self):
        if self.key_mode.get() == "auto":
            self.key_entry.config(state=tk.DISABLED,
                                  bg="#ADD8E6", fg=FG_MONO)
            self._generate_key()
        else:
            self.key_entry.config(state=tk.NORMAL,
                                  bg=MONO_BG, fg=FG_MONO)

    def _generate_key(self):
        key = generate_random_key()
        self.key_entry.config(state=tk.NORMAL)
        self._set_text(self.key_entry, key)
        if self.key_mode.get() == "auto":
            self.key_entry.config(state=tk.DISABLED,
                                  bg="#ADD8E6", fg=FG_MONO)

    def _reset_key(self):
        self.key_entry.config(state=tk.NORMAL, bg=MONO_BG)
        self.key_entry.delete("1.0", tk.END)
        self.key_mode.set("manual")
    
    def _save_key_file(self):
        key = self._get_key()
        if not key:
            messagebox.showwarning("Cảnh báo", "Chưa có khóa để lưu!")
            return
        if not is_valid_hex(key):
            messagebox.showerror("Lỗi", "Khóa phải là 16 ký tự HEX hợp lệ (0-9, A-F)!")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(key)
            messagebox.showinfo("Thông báo", f"Đã lưu khóa vào:\n{path}")

    def _open_key_file(self):
        path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
        if path:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read().strip().upper()
            self.key_entry.config(state=tk.NORMAL)
            self._set_text(self.key_entry, content)
            if self.key_mode.get() == "auto":
                self.key_mode.set("manual")
                self.key_entry.config(bg=MONO_BG, fg=FG_MONO)

    def _copy_key_to_decrypt(self):
        key = self._get_key()
        if not is_valid_hex(key):
            messagebox.showerror("Lỗi", "Khóa phải là 16 ký tự HEX hợp lệ (0-9, A-F)!")
            return
        self._set_text(self.dec_key_display, key)
        messagebox.showinfo("Thông báo", "Đã chuyển khóa sang vùng giải mã!")

    def _open_plain_file(self):
        path = filedialog.askopenfilename(filetypes=[
            ("Tất cả hỗ trợ", "*.txt *.docx *.pdf"),
            ("Text files", "*.txt"),
            ("Word files", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")])
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".docx":
                content = read_docx_file(path)
            elif ext == ".pdf":
                content = read_pdf_file(path)
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
            self._set_text(self.plain_text, content)
        except Exception as e:
            messagebox.showerror("Lỗi đọc file", str(e))

    def _save_plain_input_file(self):
        plain = self.plain_text.get("1.0", tk.END).rstrip("\n")
        if not plain:
            messagebox.showwarning("Cảnh báo", "Chưa có bản rõ để lưu!")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt",
                                            filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")])
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".docx":
                write_docx_file(path, plain)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(plain)
            messagebox.showinfo("Thông báo", f"Đã lưu bản rõ vào:\n{path}")
        except Exception as e:
            messagebox.showerror("Lỗi lưu file", str(e))

    def _encrypt(self):
        plaintext = self.plain_text.get("1.0", tk.END).rstrip("\n")
        if not plaintext:
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập bản rõ!")
            return
        key = self._get_key()
        if not is_valid_hex(key):
            messagebox.showerror("Lỗi", "Khóa phải là 16 ký tự HEX hợp lệ (0-9, A-F)!")
            return
        try:
            cipher_hex = des_encrypt(plaintext, key)
            self.last_plaintext = plaintext
            self.last_ciphertext_hex = cipher_hex
            self.last_key_used = key
            self._set_readonly(self.cipher_text, cipher_hex)
            messagebox.showinfo("Thông báo", "✅ Mã hóa thành công!")
        except Exception as e:
            messagebox.showerror("Lỗi mã hóa", str(e))

    def _transfer_cipher(self):
        cipher = self.cipher_text.get("1.0", tk.END).strip()
        if not cipher:
            messagebox.showwarning("Cảnh báo", "Chưa có bản mã để chuyển!")
            return
        key = self._get_key()
        if not is_valid_hex(key):
            messagebox.showerror("Lỗi", "Khóa phải là 16 ký tự HEX hợp lệ (0-9, A-F)!")
            return
        self._set_text(self.dec_input, cipher)
        self._set_text(self.dec_key_display, key)
        messagebox.showinfo("Thông báo", "Đã chuyển bản mã và khóa sang vùng giải mã!")

    def _save_cipher_file(self):
        cipher = self.cipher_text.get("1.0", tk.END).strip()
        if not cipher:
            messagebox.showwarning("Cảnh báo", "Chưa có bản mã để lưu!")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt",
                                            filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")])
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".docx":
                write_docx_file(path, cipher)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(cipher)
            messagebox.showinfo("Thông báo", f"Đã lưu bản mã vào:\n{path}")
        except Exception as e:
            messagebox.showerror("Lỗi lưu file", str(e))

    def _open_cipher_file(self):
        path = filedialog.askopenfilename(filetypes=[
            ("Tất cả hỗ trợ", "*.txt *.docx"),
            ("Text files", "*.txt"),
            ("Word files", "*.docx"),
            ("All files", "*.*")])
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".docx":
                content = read_docx_file(path)
            else:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
            self._set_text(self.dec_input, content.strip().upper())
        except Exception as e:
            messagebox.showerror("Lỗi đọc file", str(e))

    def _decrypt(self):
        cipher_hex = self.dec_input.get("1.0", tk.END).strip().upper()
        if not cipher_hex:
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập bản mã!")
            return
        cipher_invalid = (len(cipher_hex) % 16 != 0
                          or not all(c in "0123456789ABCDEF" for c in cipher_hex))
        key = self._get_dec_key()
        key_invalid = not is_valid_hex(key)

        if cipher_invalid and key_invalid:
            messagebox.showerror("Lỗi", "Bản mã và khóa đã bị thay đổi!")
            return
        elif cipher_invalid:
            messagebox.showerror("Lỗi", "Bản mã đã bị thay đổi!")
            return
        elif key_invalid:
            messagebox.showerror("Lỗi", "Khóa đã bị thay đổi!")
            return

        if self.last_ciphertext_hex and self.last_key_used:
            cipher_changed = (cipher_hex != self.last_ciphertext_hex)
            key_changed = (key != self.last_key_used)
            if cipher_changed and key_changed:
                messagebox.showerror("Lỗi", "Bản mã và khóa đã bị thay đổi!")
                return
            elif cipher_changed:
                messagebox.showerror("Lỗi", "Bản mã đã bị thay đổi!")
                return
            elif key_changed:
                messagebox.showerror("Lỗi", "Khóa đã bị thay đổi!")
                return

        try:
            plaintext = des_decrypt(cipher_hex, key)
            self._set_readonly(self.dec_output, plaintext)
            messagebox.showinfo("Thông báo", "✅ Giải mã thành công!")
        except Exception as e:
            messagebox.showerror("Lỗi giải mã", str(e))

    def _save_plain_file(self):
        plain = self.dec_output.get("1.0", tk.END).strip()
        if not plain:
            messagebox.showwarning("Cảnh báo", "Chưa có bản rõ để lưu!")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt",
                                            filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")])
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".docx":
                write_docx_file(path, plain)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(plain)
            messagebox.showinfo("Thông báo", f"Đã lưu bản rõ vào:\n{path}")
        except Exception as e:
            messagebox.showerror("Lỗi lưu file", str(e))

    def _verify_result(self):
        decrypted = self.dec_output.get("1.0", tk.END).strip()
        original = self.plain_text.get("1.0", tk.END).rstrip("\n")
        if not decrypted:
            messagebox.showwarning("Cảnh báo", "Chưa có kết quả giải mã!")
            return
        if not original:
            messagebox.showinfo("Kiểm tra", "Không có bản rõ gốc để so sánh.\n"
                                            "Hãy nhập bản rõ ở vùng Mã hóa trước.")
            return
        if decrypted == original:
            messagebox.showinfo("Kết quả kiểm tra",
                                "Bản rõ sau giải mã khớp với bản rõ gốc!\n\n"
                                f"Bản rõ: {original[:80]}{'...' if len(original)>80 else ''}")
        else:
            messagebox.showwarning("Kết quả kiểm tra",
                                   "Bản rõ sau giải mã không khớp với bản rõ gốc!\n\n"
                                   f"Gốc:       {original[:50]}\n"
                                   f"Giải mã:   {decrypted[:50]}")

if __name__ == "__main__":
    root = tk.Tk()
    app = DESApp(root)
    root.mainloop()